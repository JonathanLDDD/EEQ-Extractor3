import os
import dropbox
import fitz
import streamlit as st
from docx import Document
from openai import OpenAI
from docx import Document as DocWriter
from io import BytesIO
from datetime import datetime
from dropbox.exceptions import AuthError, ApiError

# ========== 配置 ==========
# ── 从 secrets 里读取（注意已改成分组）
DBX_APP_KEY = st.secrets["dropbox"]["app_key"]
DBX_APP_SECRET = st.secrets["dropbox"]["app_secret"]
DBX_REFRESH_TOKEN = st.secrets["dropbox"]["refresh_token"]

OPENAI_API_KEY = st.secrets["openai"]["api_key"]
TARGET_FOLDER = "/Category 1 - EEQ Preparation"
PROMPT_PATH = "Prompt.txt"
MODEL_NAME = "gpt-4-1106-preview"
client = OpenAI(api_key=OPENAI_API_KEY)

# 统一创建 Dropbox 客户端（自动续期）
def get_dbx() -> dropbox.Dropbox:
    return dropbox.Dropbox(
        oauth2_refresh_token=DBX_REFRESH_TOKEN,
        app_key=DBX_APP_KEY,
        app_secret=DBX_APP_SECRET,
        timeout=60
    )

# ========== Dropbox 工具 ==========
def list_dropbox_files(folder_path):
    try:
        dbx = get_dbx()
        result = dbx.files_list_folder(folder_path)
        files = []
        while True:
            for entry in result.entries:
                if isinstance(entry, dropbox.files.FileMetadata) and (
                    entry.name.lower().endswith(".pdf") or entry.name.lower().endswith(".docx")
                ):
                    files.append((entry.name, entry.path_lower))
            if result.has_more:
                result = dbx.files_list_folder_continue(result.cursor)
            else:
                break
        return files
    except AuthError:
        st.error("Dropbox 认证失败（可能 refresh token 或权限无效）。请在 secrets 中更新 app_key/app_secret/refresh_token。")
        st.stop()
    except ApiError as e:
        st.error(f"Dropbox API 错误：{e}")
        st.stop()
    except Exception as e:
        st.error(f"列目录失败：{e}")
        st.stop()

def download_dropbox_files(files, selected_names):
    try:
        dbx = get_dbx()
        downloaded = []
        selected = set(selected_names)
        for name, path in files:
            if name in selected:
                metadata, res = dbx.files_download(path)
                downloaded.append((name, res.content))
        return downloaded
    except AuthError:
        st.error("Dropbox 认证失败（下载阶段）。请检查 refresh token 与应用权限。")
        st.stop()
    except ApiError as e:
        st.error(f"Dropbox API 错误：{e}")
        st.stop()
    except Exception as e:
        st.error(f"下载失败：{e}")
        st.stop()

# ========== 文件 & GPT ==========
def read_txt(path):
    if not os.path.exists(path):
        st.error(f"Prompt 文件不存在: {path}")
        st.stop()
    with open(path, "r", encoding="utf-8") as f:
        return f.read()

def extract_text_from_docx_bytes(file_bytes):
    with BytesIO(file_bytes) as f:
        doc = Document(f)
        return "\n".join(p.text for p in doc.paragraphs)

def extract_text_from_pdf_bytes(file_bytes):
    with BytesIO(file_bytes) as f:
        doc = fitz.open(stream=f.read(), filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        return text

def ask_gpt(prompt):
    response = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[
            {"role": "system", "content": "You are an assistant helping QA Commons extract and format EEQs."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.2,
    )
    return response.choices[0].message.content

def write_output_to_word(results):
    doc = DocWriter()
    for filename, output in results:
        doc.add_heading(f"EEQ Output for {filename}", level=1)
        doc.add_paragraph(output)
        doc.add_page_break()
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ========== Streamlit 主逻辑 ==========
st.title("EEQ Syllabus Processor")

# 1. 列出 Dropbox 文件
st.subheader("Step 1: Select files to process")
files = list_dropbox_files(TARGET_FOLDER)
if not files:
    st.error("No files found. Please check the Dropbox folder or permissions.")
    st.stop()

file_names = [f[0] for f in files]
selected_files = st.multiselect("Select files to process:", file_names)

# 2. 开始处理
if st.button("Start Processing") and selected_files:
    st.info("Downloading and processing files, please wait...")
    downloaded_files = download_dropbox_files(files, selected_files)
    base_prompt = read_txt(PROMPT_PATH)
    results = []
    progress = st.progress(0)

    for idx, (name, content) in enumerate(downloaded_files):
        st.write(f"Processing: {name}")
        if name.lower().endswith(".docx"):
            syllabus_text = extract_text_from_docx_bytes(content)
        else:
            syllabus_text = extract_text_from_pdf_bytes(content)

        full_prompt = (
            base_prompt.strip()
            + "\n\n---\n\nNow analyze the following course syllabus and generate the EEQ extraction "
            + "in the same format as above.\n\nCourse Syllabus:\n\n"
            + syllabus_text
        )
        output = ask_gpt(full_prompt)
        results.append((name, output))
        progress.progress((idx + 1) / len(downloaded_files))

    # 3. 生成 Word
    buffer = write_output_to_word(results)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_filename = f"EEQ_Output_{timestamp}.docx"
    st.success("Processing completed!")

    # 4. 下载按钮
    st.download_button(
        label="Download Results",
        data=buffer,
        file_name=output_filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
