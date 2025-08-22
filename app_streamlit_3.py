import os
from io import BytesIO
from datetime import datetime

import streamlit as st
import dropbox
from dropbox.exceptions import AuthError, ApiError

import fitz  # PyMuPDF
from docx import Document as DocReader
from docx import Document as DocWriter
from openai import OpenAI

# ===== Secrets =====
DBX_APP_KEY = st.secrets["dropbox"]["app_key"]
DBX_APP_SECRET = st.secrets["dropbox"]["app_secret"]
DBX_REFRESH_TOKEN = st.secrets["dropbox"]["refresh_token"]
OPENAI_API_KEY = st.secrets["openai"]["api_key"]

# 初始展示路径仍然用 "/"，对用户友好；内部会自动转成 "" 再调 API
DEFAULT_START_FOLDER = "/"

PROMPT_PATH = "Prompt.txt"
MODEL_NAME = "gpt-4-1106-preview"  # 如报模型不可用可换 "gpt-4o-mini"
client = OpenAI(api_key=OPENAI_API_KEY)

# ===== Dropbox =====
def get_dbx() -> dropbox.Dropbox:
    base = dropbox.Dropbox(
        oauth2_refresh_token=DBX_REFRESH_TOKEN,
        app_key=DBX_APP_KEY,
        app_secret=DBX_APP_SECRET,
        timeout=60,
    )
    # 绑定到账号的根命名空间，避免在团队空间/共享空间下看不到子层级
    try:
        acc = base.users_get_current_account()
        ns = getattr(acc.root_info, "root_namespace_id", None)
        if ns:
            return base.with_path_root(dropbox.common.PathRoot.namespace_id(ns))
    except Exception:
        pass
    return base

# 把显示路径转换为 Dropbox API 路径
def to_api_path(display_path: str) -> str:
    """用户看到的根目录是 '/', 但 Dropbox API 需要传 ''。其他路径保持原样。"""
    if display_path.strip() == "/":
        return ""
    return display_path

def list_dropbox_folders(folder_path: str = "/"):
    try:
        dbx = get_dbx()
        api_path = to_api_path(folder_path)

        result = dbx.files_list_folder(
            api_path,
            recursive=False,
            include_mounted_folders=True,
            include_non_downloadable_files=True,
        )

        folders = []
        while True:
            for entry in result.entries:
                if isinstance(entry, dropbox.files.FolderMetadata):
                    folders.append((entry.name, entry.path_lower))
                elif hasattr(entry, "is_downloadable") and (not entry.is_downloadable):
                    if hasattr(entry, "path_lower") and hasattr(entry, "name"):
                        folders.append((entry.name, entry.path_lower))
            if result.has_more:
                result = dbx.files_list_folder_continue(result.cursor)
            else:
                break

        folders.sort(key=lambda x: x[0].lower())
        return folders
    except AuthError:
        st.error("Dropbox 认证失败，请检查 refresh token 与应用权限。"); st.stop()
    except ApiError as e:
        st.error(f"Dropbox API 错误：{e}"); st.stop()
    except Exception as e:
        st.error(f"读取文件夹失败：{e}"); st.stop()

def list_dropbox_files(folder_path: str):
    try:
        dbx = get_dbx()
        api_path = to_api_path(folder_path)
        result = dbx.files_list_folder(
            api_path,
            recursive=False,
            include_mounted_folders=True,
            include_non_downloadable_files=True,
        )

        files = []
        while True:
            for entry in result.entries:
                if isinstance(entry, dropbox.files.FileMetadata):
                    name_l = entry.name.lower()
                    if name_l.endswith(".pdf") or name_l.endswith(".docx"):
                        files.append((entry.name, entry.path_lower))
            if result.has_more:
                result = dbx.files_list_folder_continue(result.cursor)
            else:
                break

        files.sort(key=lambda x: x[0].lower())
        return files
    except AuthError:
        st.error("Dropbox 认证失败，请检查 refresh token 与应用权限。"); st.stop()
    except ApiError as e:
        st.error(f"Dropbox API 错误：{e}"); st.stop()
    except Exception as e:
        st.error(f"列出文件失败：{e}"); st.stop()

def download_dropbox_files(files, selected_names):
    try:
        dbx = get_dbx()
        selected = set(selected_names)
        out = []
        for name, path in files:
            if name in selected:
                _, res = dbx.files_download(path)
                out.append((name, res.content))
        return out
    except AuthError:
        st.error("Dropbox 认证失败（下载阶段）。"); st.stop()
    except ApiError as e:
        st.error(f"Dropbox API 错误：{e}"); st.stop()
    except Exception as e:
        st.error(f"下载失败：{e}"); st.stop()

# ===== 本地文件处理 =====
def extract_text_from_docx_bytes(file_bytes):
    with BytesIO(file_bytes) as f:
        doc = DocReader(f)
        return "\n".join(p.text for p in doc.paragraphs)

def extract_text_from_pdf_bytes(file_bytes):
    with BytesIO(file_bytes) as f:
        doc = fitz.open(stream=f.read(), filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    return text

def ask_gpt(prompt: str) -> str:
    resp = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[
            {"role": "system", "content": "You are an assistant helping QA Commons extract and format EEQs."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.2,
    )
    return resp.choices[0].message.content

def read_txt(path):
    if not os.path.exists(path):
        st.error(f"Prompt 文件不存在：{path}"); st.stop()
    with open(path, "r", encoding="utf-8") as f:
        return f.read()

def write_output_to_word(results):
    doc = DocWriter()
    for filename, output in results:
        doc.add_heading(f"EEQ Output for {filename}", level=1)
        doc.add_paragraph(output)
        doc.add_page_break()
    buf = BytesIO()
    doc.save(buf); buf.seek(0)
    return buf

# ===== UI =====
st.title("EEQ Syllabus Processor")

# Step 0: 选择 Dropbox 文件夹
st.subheader("Step 0: Choose a Dropbox folder")
if "cwd" not in st.session_state:
    st.session_state.cwd = DEFAULT_START_FOLDER

col1, col2 = st.columns([3, 1])
with col1:
    new_path = st.text_input("Current path", value=st.session_state.cwd, key="cwd_input")
with col2:
    if st.button("Go"):
        st.session_state.cwd = new_path or "/"
        st.rerun()

folders = list_dropbox_folders(st.session_state.cwd)
options = [".. (parent directory)"] + [f"{name} — {path}" for name, path in folders]
choice = st.selectbox("Folders in this path:", options)

go_col1, _ = st.columns([1, 4])
with go_col1:
    if st.button("Open selected"):
        if choice.startswith(".."):
            p = st.session_state.cwd.rstrip("/")
            parent = "/" if p in ("", "/") else "/".join(p.split("/")[:-1])
            if parent == "":
                parent = "/"
            st.session_state.cwd = parent
        else:
            sel_path = choice.split("—")[-1].strip()
            st.session_state.cwd = sel_path
        st.rerun()

st.caption("Tip: Type a path (e.g., /Dept/Hist/Syllabi) or pick a subfolder then click 'Open selected'.")

# Step 1A: 选择当前 Dropbox 文件夹中的大纲（不递归）
st.subheader("Step 1A: Select syllabi from the current Dropbox folder")
files = list_dropbox_files(st.session_state.cwd)
dropbox_file_names = [f[0] for f in files]
selected_dropbox_files = st.multiselect("Pick Dropbox files to process:", dropbox_file_names)

# Step 1B: 或者上传本地文件（支持多选）
st.subheader("Step 1B: Or upload local files")
uploaded_files = st.file_uploader(
    "Upload one or more files (.pdf or .docx)",
    type=["pdf", "docx"],
    accept_multiple_files=True,
    help="You can combine uploads with Dropbox selections above.",
)

# Step 2: 处理
st.subheader("Step 2: Run")
run_clicked = st.button("Start Processing")

if run_clicked:
    # 1) 合并两种来源的文件到统一结构：[(name, bytes), ...]
    combined_inputs = []

    # 从 Dropbox 下载被选中的文件
    if selected_dropbox_files:
        st.info("Downloading from Dropbox...")
        combined_inputs.extend(download_dropbox_files(files, selected_dropbox_files))

    # 加入本地上传的文件
    if uploaded_files:
        for up in uploaded_files:
            combined_inputs.append((up.name, up.read()))

    if not combined_inputs:
        st.warning("Please select at least one Dropbox file or upload local files.")
        st.stop()

    # 2) 读取 Prompt
    base_prompt = read_txt(PROMPT_PATH)

    # 3) 逐个处理
    st.info("Processing files...")
    results = []
    progress = st.progress(0)

    for idx, (name, content) in enumerate(combined_inputs):
        st.write(f"Processing: {name}")
        if name.lower().endswith(".docx"):
            syllabus_text = extract_text_from_docx_bytes(content)
        else:
            syllabus_text = extract_text_from_pdf_bytes(content)

        full_prompt = (
            base_prompt.strip()
            + "\n\n---\n\nNow analyze the following course syllabus and generate the EEQ extraction "
            + "in the same format as above.\n\nCourse Syllabus:\n\n"
            + syllabus_text
        )
        output = ask_gpt(full_prompt)
        results.append((name, output))
        progress.progress((idx + 1) / len(combined_inputs))

    # 4) 打包下载
    buf = write_output_to_word(results)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_name = f"EEQ_Output_{ts}.docx"
    st.success("Processing completed!")
    st.download_button(
        label="Download Results",
        data=buf,
        file_name=out_name,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
